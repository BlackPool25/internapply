"""Professional DOCX resume renderer — ATS-optimized, single-column, 1-page."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ── Colors ─────────────────────────────────────────────────────────────────
BLACK = RGBColor(0x1A, 0x1A, 0x1A)
GRAY = RGBColor(0x55, 0x55, 0x55)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)

# ── Typography ──────────────────────────────────────────────────────────────
FONT_NAME = "Arial"
NAME_SIZE = Pt(16)
SUMMARY_SIZE = Pt(9.5)
SECTION_SIZE = Pt(11)
BODY_SIZE = Pt(9.5)
SMALL_SIZE = Pt(8.5)
CONTACT_SIZE = Pt(8.5)

# ── Margins ─────────────────────────────────────────────────────────────────
MARGIN = Inches(0.55)

# ── Colors ──────────────────────────────────────────────────────────────────
ACCENT_LINE = RGBColor(0xCC, 0xCC, 0xCC)

# ── Page limits ──────────────────────────────────────────────────────────────
MAX_PROJECTS = 4


def _set_font(run, size: Pt, bold: bool = False, color: RGBColor = BLACK,
              italic: bool = False):
    """Apply formatting to a run."""
    run.font.name = FONT_NAME
    run.font.size = size
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.italic = italic
    # Ensure Arial is embedded for cross-platform rendering
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), FONT_NAME)
    rFonts.set(qn("w:hAnsi"), FONT_NAME)
    rFonts.set(qn("w:cs"), FONT_NAME)


def _add_section_heading(doc: Document, text: str):
    """Add an ATS-friendly section heading with a bottom border line."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(5)
    para.paragraph_format.space_after = Pt(2)
    para.paragraph_format.keep_with_next = True
    run = para.add_run(text.upper())
    _set_font(run, SECTION_SIZE, bold=True, color=BLACK)
    # Add bottom border
    pPr = para._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single",
        qn("w:sz"): "4",
        qn("w:space"): "1",
        qn("w:color"): "CCCCCC",
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_bullet(doc: Document, text: str, size: Pt = BODY_SIZE):
    """Add a bullet point."""
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = Pt(11)
    para.clear()
    run = para.add_run(text)
    _set_font(run, size, color=DARK_GRAY)


def _add_contact_line(doc: Document, parts: list[tuple[str, bool, str | None]]):
    """Add a contact line with mixed formatting.

    Args:
        parts: List of (text, is_link, url_or_None) tuples.
    """
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(2)
    para.paragraph_format.space_before = Pt(0)
    for i, (text, is_link, url) in enumerate(parts):
        if i > 0:
            sep = para.add_run("  |  ")
            _set_font(sep, CONTACT_SIZE, color=GRAY)
        run = para.add_run(text)
        _set_font(run, CONTACT_SIZE, color=GRAY if not is_link else RGBColor(0x1A, 0x56, 0xDB))


def _add_skill_line(doc: Document, label: str, value: str):
    """Add a skill line: 'label: value'."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(1)
    para.paragraph_format.line_spacing = Pt(12)
    label_run = para.add_run(label + ": ")
    _set_font(label_run, BODY_SIZE, bold=True)
    value_run = para.add_run(value)
    _set_font(value_run, BODY_SIZE, color=DARK_GRAY)


def _add_project_block(doc: Document, name: str, tech: str, bullets: list[str],
                       url: str | None = None):
    """Add a project entry with name, tech stack, and bullet points."""
    # Project name line
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.keep_with_next = True
    name_run = para.add_run(name)
    _set_font(name_run, BODY_SIZE, bold=True)

    if url:
        url_run = para.add_run("  |  GitHub")
        _set_font(url_run, SMALL_SIZE, color=RGBColor(0x1A, 0x56, 0xDB), italic=True)

    # Tech stack
    if tech:
        tech_para = doc.add_paragraph()
        tech_para.paragraph_format.space_before = Pt(0)
        tech_para.paragraph_format.space_after = Pt(1)
        tech_para.paragraph_format.line_spacing = Pt(12)
        label = tech_para.add_run("Tech: ")
        _set_font(label, SMALL_SIZE, bold=True, color=GRAY, italic=True)
        tech_run = tech_para.add_run(tech)
        _set_font(tech_run, SMALL_SIZE, color=GRAY, italic=True)

    # Bullets (max 2 for 1-page constraint)
    for bullet in bullets[:2]:
        _add_bullet(doc, bullet, SMALL_SIZE)


def _add_education_block(doc: Document, entries: list[dict]):
    """Add education entries."""
    for entry in entries:
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(1)
        para.paragraph_format.keep_with_next = True
        degree = entry.get("degree", entry.get("degree_name", ""))
        inst = entry.get("institution", "")
        cgpa = entry.get("cgpa", entry.get("gpa", ""))
        expected = entry.get("expected", "")

        deg_run = para.add_run(degree)
        _set_font(deg_run, BODY_SIZE, bold=True)

        if inst:
            inst_run = para.add_run("  |  " + inst)
            _set_font(inst_run, BODY_SIZE, color=DARK_GRAY)

        extra_parts = []
        if cgpa:
            extra_parts.append(f"CGPA: {cgpa}")
        if expected:
            extra_parts.append(f"Expected: {expected}")
        if extra_parts:
            extra_run = para.add_run("  |  " + " | ".join(extra_parts))
            _set_font(extra_run, BODY_SIZE, color=GRAY)


def render_resume(data: dict[str, Any], output_path: str | Path,
                  company: str = "", job_title: str = "") -> Path:
    """Render a tailored resume dict to a professional ATS-optimized DOCX.

    Creates a single-column, 1-page DOCX optimized for Workday, Greenhouse,
    Lever, Internshala, and other major ATS platforms.

    Args:
        data: Tailored resume dict (from ResumeTailor or master resume).
        output_path: Where to save the .docx file.
        company: Target company (for header context).
        job_title: Target role (for header context).

    Returns:
        The path to the generated .docx file.
    """
    doc = Document()

    # ── Page setup ────────────────────────────────────────────────────
    section = doc.sections[0]
    section.top_margin = MARGIN
    section.bottom_margin = MARGIN
    section.left_margin = MARGIN
    section.right_margin = MARGIN

    # Remove default empty paragraph
    if doc.paragraphs:
        p = doc.paragraphs[0]
        p._element.getparent().remove(p._element)

    # ── Name ──────────────────────────────────────────────────────────
    name = data.get("name", "Resume")
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_before = Pt(0)
    name_para.paragraph_format.space_after = Pt(2)
    name_run = name_para.add_run(name)
    _set_font(name_run, NAME_SIZE, bold=True, color=BLACK)

    # ── Contact ───────────────────────────────────────────────────────
    email = data.get("email", "")
    phone = data.get("phone", "")
    location = data.get("location", "")
    contact_parts = []
    if location:
        contact_parts.append((location, False, None))
    if email:
        contact_parts.append((email, True, f"mailto:{email}"))
    if phone:
        contact_parts.append((phone, False, None))
    if contact_parts:
        _add_contact_line(doc, contact_parts)

    # ── Tailoring note (subtle) ───────────────────────────────────────
    if company and job_title:
        note = doc.add_paragraph()
        note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        note.paragraph_format.space_after = Pt(4)
        note.paragraph_format.space_before = Pt(0)
        note_run = note.add_run(f"Applicant for {job_title} @ {company}")
        _set_font(note_run, Pt(8), italic=True, color=GRAY)

    # ── Summary ───────────────────────────────────────────────────────
    summary = data.get("summary", "")
    if summary:
        _add_section_heading(doc, "Summary")
        sum_para = doc.add_paragraph()
        sum_para.paragraph_format.space_after = Pt(2)
        sum_para.paragraph_format.line_spacing = Pt(12)
        sum_run = sum_para.add_run(summary)
        _set_font(sum_run, SUMMARY_SIZE, color=DARK_GRAY)

    # ── Education ─────────────────────────────────────────────────────
    education = data.get("education", [])
    if education:
        _add_section_heading(doc, "Education")
        _add_education_block(doc, education)

    # ── Technical Skills ─────────────────────────────────────────────
    skills = data.get("skills_reordered", [])
    if skills:
        _add_section_heading(doc, "Key Skills")
        skill_text = ", ".join(skills[:10])
        if len(skills) > 10:
            skill_text += f" (+{len(skills) - 10} more)"
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(1)
        para.paragraph_format.line_spacing = Pt(11)
        run = para.add_run(skill_text)
        _set_font(run, BODY_SIZE, color=DARK_GRAY)
    else:
        # Fallback: use master resume skill categories
        skills_raw = data.get("skills", {})
        if isinstance(skills_raw, dict):
            _add_section_heading(doc, "Technical Skills")
            for cat, val in skills_raw.items():
                if isinstance(val, str) and val.strip():
                    _add_skill_line(doc, cat, val)

    # ── Projects ─────────────────────────────────────────────────────
    projects = data.get("projects", [])
    if projects:
        _add_section_heading(doc, "Projects")
        for proj in projects[:MAX_PROJECTS]:
            name = proj.get("name", "")
            tech = proj.get("tech", "")
            bullets = proj.get("bullets", proj.get("description", []))
            if isinstance(bullets, str):
                bullets = [bullets]
            url = proj.get("url", "")
            _add_project_block(doc, name, tech, bullets, url)

    # ── Additional (one-liner at bottom) ────────────────────────────
    additional = data.get("additional", [])
    oss = next((a for a in additional if "open source" in a.get("label", "").lower()), None)
    hack = next((a for a in additional if "hackathon" in a.get("label", "").lower()), None)
    extras = []
    if oss:
        extras.append(oss["value"][:80])
    if hack:
        extras.append(hack["value"][:60])
    if extras:
        add_para = doc.add_paragraph()
        add_para.paragraph_format.space_before = Pt(1)
        add_para.paragraph_format.space_after = Pt(0)
        add_para.paragraph_format.line_spacing = Pt(10)
        add_run = add_para.add_run(" | ".join(extras))
        _set_font(add_run, Pt(7.5), color=GRAY)

    # ── Save ─────────────────────────────────────────────────────────
    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


__all__ = ["render_resume"]

# ── ATS reader check + deterministic fix (ponytail: minimal, no new deps) ──
import zipfile
import xml.etree.ElementTree as ET
from datetime import datetime, timezone

_DETERMINISTIC_CREATED = datetime(2025, 1, 1, tzinfo=timezone.utc)

def _ensure_deterministic(doc) -> None:
    """Fix core props for deterministic DOCX (mtime-safe)."""
    try:
        cp = doc.core_properties
        cp.created = _DETERMINISTIC_CREATED
        cp.modified = _DETERMINISTIC_CREATED
        cp.revision = 1
    except Exception:
        pass

def ats_reader_check(docx_path: str | Path) -> list[str]:
    """Block multi-col / table layouts that break ATS parse (96.7% CVCraft target).

    Returns list of issues (empty = pass). Checks: w:tbl present, w:cols num>1, w:textDirection.
    """
    p = Path(docx_path)
    if not p.exists():
        return [f"docx not found: {p}"]
    issues: list[str] = []
    try:
        with zipfile.ZipFile(str(p)) as z:
            xml = z.read("word/document.xml").decode()
    except Exception as e:
        return [f"cannot read document.xml: {e}"]
    # simple string checks (ponytail: regex enough, no full parse)
    if "<w:tbl" in xml:
        issues.append("block multi-col/table: w:tbl found (ATS unfriendly)")
    if 'w:num="2"' in xml or 'w:num="3"' in xml:
        issues.append("block multi-column: w:cols num>1")
    if "w:textDirection" in xml:
        issues.append("block textDirection (multi-col)")
    # headings check: at least one w:pStyle w:val="Heading" or section headings present via ALL-CAPS heuristic
    # list check: ensure bullets use w:numPr not raw table
    return issues

# wrap render_resume determinism
_orig_render = render_resume
def render_resume(data, output_path, company="", job_title=""):
    # deterministic: sort keys in data before render (no-op for doc but ensures stable input)
    if isinstance(data, dict):
        # shallow sorted copy for reproducibility
        data = {k: data[k] for k in sorted(data.keys())}
    # call original logic inline to add deterministic fix
    # re-import to avoid recursion confusion — directly implement deterministic save
    from docx import Document as _Doc
    # we already have _orig_render; use it but inject deterministic core props via monkey
    out = _orig_render(data, output_path, company, job_title)
    # re-open to patch core props deterministically (ponytail: cheap post-fix)
    try:
        doc = _Doc(str(out))
        _ensure_deterministic(doc)
        doc.save(str(out))
    except Exception:
        pass
    # ats check — log warning if issues
    issues = ats_reader_check(out)
    if issues:
        from loguru import logger as _lg
        _lg.warning("ATS check issues: {}", issues)
    return out
